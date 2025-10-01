# src/utils/file_handler.py
import os
import tempfile
import logging
import mimetypes
import time
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, BinaryIO
import PyPDF2
import docx
import pydicom
from PIL import Image
import cv2
import numpy as np
import streamlit as st

class FileHandler:
    """Handles file operations for the medical assistant"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.temp_dir = tempfile.gettempdir()
        
        # Supported file types
        self.supported_image_types = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.dcm'}
        self.supported_document_types = {'.pdf', '.txt', '.doc', '.docx'}
        self.supported_types = self.supported_image_types | self.supported_document_types
        
        # Maximum file sizes (in bytes)
        self.max_image_size = 50 * 1024 * 1024  # 50MB
        self.max_document_size = 100 * 1024 * 1024  # 100MB
    
    def validate_file(self, uploaded_file) -> Dict[str, Any]:
        """Validate uploaded file"""
        try:
            file_info = {
                'valid': False,
                'filename': uploaded_file.name,
                'size': 0,
                'type': 'unknown',
                'mime_type': '',
                'errors': []
            }
            
            # Get file extension
            file_ext = Path(uploaded_file.name).suffix.lower()
            
            # Check if file type is supported
            if file_ext not in self.supported_types:
                file_info['errors'].append(f"Unsupported file type: {file_ext}")
                return file_info
            
            # Get file size
            if hasattr(uploaded_file, 'size'):
                file_size = uploaded_file.size
            else:
                # For streamlit uploaded files, read to get size
                current_position = uploaded_file.tell()
                uploaded_file.seek(0, 2)  # Seek to end
                file_size = uploaded_file.tell()
                uploaded_file.seek(current_position)  # Restore position
            
            file_info['size'] = file_size
            
            # Determine file type
            if file_ext in self.supported_image_types:
                file_info['type'] = 'image'
                max_size = self.max_image_size
            elif file_ext in self.supported_document_types:
                file_info['type'] = 'document'
                max_size = self.max_document_size
            else:
                file_info['errors'].append("Unknown file category")
                return file_info
            
            # Check file size
            if file_size > max_size:
                file_info['errors'].append(f"File too large: {file_size / (1024*1024):.1f}MB (max: {max_size / (1024*1024):.1f}MB)")
                return file_info
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(uploaded_file.name)
            file_info['mime_type'] = mime_type or 'application/octet-stream'
            
            # Additional validation based on file type
            if file_ext == '.dcm':
                if not self._validate_dicom_file(uploaded_file):
                    file_info['errors'].append("Invalid DICOM file format")
                    return file_info
            
            # If no errors, file is valid
            if not file_info['errors']:
                file_info['valid'] = True
            
            return file_info
            
        except Exception as e:
            self.logger.error(f"Error validating file: {e}")
            return {
                'valid': False,
                'filename': getattr(uploaded_file, 'name', 'unknown'),
                'errors': [f"Validation error: {str(e)}"]
            }
    
    def save_temp_file(self, uploaded_file) -> str:
        """Save uploaded file to temporary location"""
        try:
            # Create unique temporary filename
            file_ext = Path(uploaded_file.name).suffix
            temp_file = tempfile.NamedTemporaryFile(
                delete=False, 
                suffix=file_ext,
                dir=self.temp_dir
            )
            
            # Write file content
            if hasattr(uploaded_file, 'read'):
                content = uploaded_file.read()
                temp_file.write(content)
                # Reset file pointer for potential future reads
                if hasattr(uploaded_file, 'seek'):
                    uploaded_file.seek(0)
            else:
                # Handle different file object types
                with open(uploaded_file, 'rb') as f:
                    content = f.read()
                    temp_file.write(content)
            
            temp_file.close()
            
            self.logger.info(f"Saved temporary file: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            self.logger.error(f"Error saving temporary file: {e}")
            raise Exception(f"Failed to save file: {str(e)}")
    
    def extract_text(self, uploaded_file) -> str:
        """Extract text from various document formats"""
        try:
            # Handle both file objects and file paths
            if isinstance(uploaded_file, str):
                # It's a file path
                file_path = uploaded_file
                file_ext = Path(file_path).suffix.lower()
            else:
                # It's a file object
                file_path = uploaded_file
                file_ext = Path(uploaded_file.name).suffix.lower()
            
            if file_ext == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported document format: {file_ext}")
                
        except Exception as e:
            self.logger.error(f"Error extracting text: {e}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_text_from_txt(self, file_input) -> str:
        """Extract text from TXT file"""
        try:
            if isinstance(file_input, str):
                # It's a file path
                with open(file_input, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                # It's a file object
                if hasattr(file_input, 'read'):
                    content = file_input.read()
                    if isinstance(content, bytes):
                        # Try different encodings
                        for encoding in ['utf-8', 'latin-1', 'cp1252']:
                            try:
                                return content.decode(encoding)
                            except UnicodeDecodeError:
                                continue
                        # If all encodings fail, use utf-8 with error handling
                        return content.decode('utf-8', errors='ignore')
                    else:
                        return str(content)
                else:
                    raise ValueError("Invalid file input type")
                    
        except Exception as e:
            self.logger.error(f"Error extracting text from TXT: {e}")
            raise
    
    def _extract_text_from_pdf(self, file_input) -> str:
        """Extract text from PDF file"""
        try:
            text_content = []
            
            if isinstance(file_input, str):
                # It's a file path
                with open(file_input, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    # Extract text from all pages
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            else:
                # It's a file object
                if hasattr(file_input, 'read'):
                    pdf_reader = PyPDF2.PdfReader(file_input)
                else:
                    with open(file_input, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF: {e}")
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def _extract_text_from_docx(self, file_input) -> str:
        """Extract text from DOCX file"""
        try:
            if isinstance(file_input, str):
                # It's a file path
                doc = docx.Document(file_input)
            else:
                # It's a file object
                if hasattr(file_input, 'read'):
                    # Save to temporary file first for docx processing
                    temp_path = self.save_temp_file(file_input)
                    doc = docx.Document(temp_path)
                    # Clean up temp file
                    os.unlink(temp_path)
                else:
                    doc = docx.Document(file_input)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX: {e}")
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def _validate_dicom_file(self, uploaded_file) -> bool:
        """Validate DICOM file format"""
        try:
            # Save to temp file for DICOM validation
            temp_path = self.save_temp_file(uploaded_file)
            
            try:
                # Try to read DICOM file
                dicom_data = pydicom.dcmread(temp_path)
                
                # Basic validation checks
                required_tags = ['StudyInstanceUID', 'SeriesInstanceUID', 'SOPInstanceUID']
                for tag in required_tags:
                    if not hasattr(dicom_data, tag):
                        return False
                
                # Check if pixel data exists for images
                if hasattr(dicom_data, 'pixel_array'):
                    pixel_array = dicom_data.pixel_array
                    if pixel_array is None or pixel_array.size == 0:
                        return False
                
                return True
                
            finally:
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            self.logger.error(f"DICOM validation error: {e}")
            return False
    
    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file"""
        try:
            metadata = {
                'filename': os.path.basename(file_path),
                'size': os.path.getsize(file_path),
                'extension': Path(file_path).suffix.lower(),
                'created': os.path.getctime(file_path),
                'modified': os.path.getmtime(file_path)
            }
            
            # Add specific metadata based on file type
            file_ext = metadata['extension']
            
            if file_ext == '.dcm':
                metadata.update(self._get_dicom_metadata(file_path))
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                metadata.update(self._get_image_metadata(file_path))
            elif file_ext == '.pdf':
                metadata.update(self._get_pdf_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error getting file metadata: {e}")
            return {'error': str(e)}
    
    def _get_dicom_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DICOM-specific metadata"""
        try:
            dicom_data = pydicom.dcmread(file_path)
            
            metadata = {
                'modality': getattr(dicom_data, 'Modality', 'Unknown'),
                'study_date': getattr(dicom_data, 'StudyDate', 'Unknown'),
                'patient_id': getattr(dicom_data, 'PatientID', '[DEIDENTIFIED]'),
                'study_description': getattr(dicom_data, 'StudyDescription', 'Unknown'),
                'series_description': getattr(dicom_data, 'SeriesDescription', 'Unknown'),
                'manufacturer': getattr(dicom_data, 'Manufacturer', 'Unknown'),
                'institution': '[DEIDENTIFIED]'  # Always de-identify institution
            }
            
            # Add image dimensions if available
            if hasattr(dicom_data, 'pixel_array'):
                pixel_array = dicom_data.pixel_array
                metadata['image_dimensions'] = {
                    'height': pixel_array.shape[0],
                    'width': pixel_array.shape[1],
                    'channels': len(pixel_array.shape)
                }
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting DICOM metadata: {e}")
            return {'dicom_error': str(e)}
    
    def _get_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract image metadata"""
        try:
            # Using PIL for basic image info
            with Image.open(file_path) as img:
                metadata = {
                    'dimensions': {
                        'width': img.width,
                        'height': img.height
                    },
                    'mode': img.mode,
                    'format': img.format
                }
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    metadata['exif'] = {
                        'orientation': exif_data.get(274, 'Unknown'),
                        'datetime': exif_data.get(306, 'Unknown'),
                        'camera_make': exif_data.get(271, 'Unknown'),
                        'camera_model': exif_data.get(272, 'Unknown')
                    }
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting image metadata: {e}")
            return {'image_error': str(e)}
    
    def _get_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted
                }
                
                # Extract document info if available
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_info.get('/Title', 'Unknown'),
                        'author': '[DEIDENTIFIED]',  # Always de-identify author
                        'subject': pdf_info.get('/Subject', 'Unknown'),
                        'creator': pdf_info.get('/Creator', 'Unknown'),
                        'creation_date': str(pdf_info.get('/CreationDate', 'Unknown'))
                    })
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF metadata: {e}")
            return {'pdf_error': str(e)}
    
    def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up old temporary files"""
        try:
            current_time = time.time()
            cleanup_count = 0
            
            temp_dir = Path(self.temp_dir)
            for file_path in temp_dir.glob("tmp*"):
                if file_path.is_file():
                    file_age_hours = (current_time - file_path.stat().st_mtime) / 3600
                    if file_age_hours > max_age_hours:
                        try:
                            file_path.unlink()
                            cleanup_count += 1
                        except Exception as e:
                            self.logger.warning(f"Could not delete temp file {file_path}: {e}")
            
            if cleanup_count > 0:
                self.logger.info(f"Cleaned up {cleanup_count} temporary files")
                
        except Exception as e:
            self.logger.error(f"Error during temp file cleanup: {e}")
    
    def convert_image_format(self, input_path: str, output_format: str = 'PNG') -> str:
        """Convert image to specified format"""
        try:
            with Image.open(input_path) as img:
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'LA') and output_format.upper() == 'JPEG':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Create output path
                input_path_obj = Path(input_path)
                output_path = input_path_obj.parent / f"{input_path_obj.stem}_converted.{output_format.lower()}"
                
                # Save in new format
                img.save(output_path, format=output_format.upper())
                
                return str(output_path)
                
        except Exception as e:
            self.logger.error(f"Error converting image format: {e}")
            raise Exception(f"Failed to convert image: {str(e)}")
    
    def resize_image(self, input_path: str, max_size: tuple = (1024, 1024)) -> str:
        """Resize image while maintaining aspect ratio"""
        try:
            with Image.open(input_path) as img:
                # Calculate new size maintaining aspect ratio
                img.thumbnail(max_size, Image.Resampling.LANCZOS)
                
                # Create output path
                input_path_obj = Path(input_path)
                output_path = input_path_obj.parent / f"{input_path_obj.stem}_resized{input_path_obj.suffix}"
                
                # Save resized image
                img.save(output_path)
                
                return str(output_path)
                
        except Exception as e:
            self.logger.error(f"Error resizing image: {e}")
            raise Exception(f"Failed to resize image: {str(e)}")
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats"""
        return {
            'images': list(self.supported_image_types),
            'documents': list(self.supported_document_types),
            'all': list(self.supported_types)
        }